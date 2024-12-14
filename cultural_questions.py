'''
This program is used to create graphs from the responses generated from a google form
Documentation limitations (hardcodes): questions range from 1-52, responses range from 1,5 and empty. 2 comment sections, 54 surveyee's
Author: Cooper Goddard
Date (init): 2022-11-01 
'''
import matplotlib.pyplot as plt
from math import floor
from PIL import Image
import csv
import docx

NUMBER_OF_QUESTIONS = 52
NUMBER_OF_RESPONSES = 11
NUMBER_OF_COMMENTS = 0
MIN_ANSWER = 1
MAX_ANSWER = 5


def get_csv_data(filename: str) -> list[list[int]] and list[str] and list[str]:
    '''
    Takes a csv file full of data (with a header) and breaks it into 
    a grid of arrays which it returns.
    '''
    with open(filename, "r", encoding="utf8") as file:
        rows = list(csv.reader(file, delimiter=','))
        headers = rows.pop(0)

        # Populate 2D Array
        grid = []
        comments = []
        for row in rows:
            row.pop(0)  # remove timestamp
            for c in range(NUMBER_OF_COMMENTS): # remove comment sections
                comments.append(row.pop(-1))  
            grid.append(row)

        # Make a new grid, one row for each question
        new_grid = []
        count = 1
        while count < NUMBER_OF_QUESTIONS:
            new_grid.append([])
            count += 1

        # swap rows and columns, adding to the new grid
        for row in range(len(grid)):
            for value in range(len(grid[row])):
                current_value = grid[row][value]
                if current_value == '':
                    new_grid[value].append(0)
                else:
                    new_grid[value].append(floor(float((current_value)))) # floor and integerize

    return new_grid, headers, comments


def sort_data(data: list[int]) -> list[int]:
    '''
    This function takes a list of numbers and totals the ammount of each number
    For example: [1,1,1,2,3,4,3] would return [3,1,2,1,0]
    '''
    result = [0] * MAX_ANSWER
    for value in data:
        # add count for the value to proper index
        if value != 0:
            result[value-1] += 1

    return result


def create_graph(header: str, data: list[int]) -> None:
    '''
    This function creates a graph based on the column provided (question num)
    using the the provided data. It saves the graph as an image, 'temp.jpg'.
    '''
    fig, ax = plt.subplots()
    bar_labels = []
    for num in range(MIN_ANSWER, MAX_ANSWER+1):
        bar_labels.append(f"{num}")
    bar_colors = ['tab:green', 'tab:blue', 'tab:red', 'tab:orange', 'tab:pink']

    ax.bar(bar_labels, data, label=bar_labels, color=bar_colors)
    ax.set_ylim(0, NUMBER_OF_RESPONSES)
    ax.bar_label(ax.containers[0], label_type='edge')
    ax.set_ylabel('Number of Responses')
    ax.set_title(header.split('.')[0])  # just the question number

    plt.savefig('temp.jpg')
    plt.close()


def make_document(grid: list[list[int]], headers: list[str], comments: list[str]) -> None:
    '''
    This function uses the data to build graphs and insert them into a
    word document (.docx)
    '''
    # create document to be populated
    document = docx.Document()
    document.add_heading("Cultural Questions Analysis", 0)

    # For each question, create a graph and upload to the document
    for question in range(len(grid)):
        # create and add graph
        document.add_heading(headers[question], 1)
        create_graph(headers[question], sort_data(
            grid[question]))  # graph saved as temp.jpg
        imageFile = Image.open('temp.jpg')
        document.add_picture('temp.jpg')
        imageFile.close()

    # add comments, placing space between each user's submission
    document.add_heading("Comments", 0)
    for comment_index in range(len(comments)):
        document.add_paragraph(comments[comment_index])

        # place spacing between user comments
        if comment_index % NUMBER_OF_COMMENTS:
            document.add_paragraph("")


    document.save('report.docx')
    return


def main() -> None:
    # Get Data From CSV
    grid, headers, comments = get_csv_data('responses.csv')
    headers.pop(0)  # remove timestamp (useless)

    # Create Graphs for each question and upload to document
    make_document(grid, headers, comments)

    return


if __name__ == "__main__":
    main()
